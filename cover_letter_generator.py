import os
import sys
import json
import datetime
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn

def clean_markdown(text):
    """Removes ALL markdown artifacts."""
    if not text: return ""
    text = str(text)
    import re
    text = re.sub(r'\[\[([^\]]+)\]\(([^\)]+)\)\]\(([^\)]+)\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\1', text)
    text = re.sub(r'\|\s*---+\s*\|', '', text)
    text = re.sub(r'^\|\s*|\s*\|$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*|\*', '', text)
    text = re.sub(r'__|_', '', text)
    lines = text.split('\n')
    cleaned = [s.lstrip('|').rstrip('|').strip() for s in lines if s.strip() and not (s.startswith('|') and '---' in s)]
    return '\n'.join(cleaned).strip()

def add_hyperlink(paragraph, url, text, bold=False, size=10, color='005b9f'):
    """Adds a clickable hyperlink."""
    if not url:
        run = paragraph.add_run(clean_markdown(text))
        run.font.size = Pt(size)
        if bold: run.bold = True
        return
    if not url.startswith('http'):
        url = 'https://' + url
        
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:tooltip'), url)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Calibri')
    rFont.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFont)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.set(qn('xml:space'), 'preserve')
    text_elem.text = clean_markdown(text)
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def generate_cover_letter(cv_data_input, letter_data_input, output_file="Anschreiben.docx"):
    # Allow passing either a file path or a dictionary for CV
    if isinstance(cv_data_input, str):
        try:
            with open(cv_data_input, 'r', encoding='utf-8') as f:
                cv_data = json.load(f)
        except Exception as e:
            print(f"❌ Error loading CV file: {e}")
            return
    else:
        cv_data = cv_data_input

    # Allow passing either a file path or a dictionary for Letter
    if isinstance(letter_data_input, str):
        try:
            with open(letter_data_input, 'r', encoding='utf-8') as f:
                letter_data = json.load(f)
        except Exception as e:
            print(f"❌ Error loading Letter file: {e}")
            return
    else:
        letter_data = letter_data_input

    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11) # Base text slightly larger for Cover Letter
    
    # 1. Page Margins (German Standards identical to CV)
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(25)
        section.right_margin = Mm(20)

    basics = cv_data.get("basics", {})
    recipient = letter_data.get("recipient", {})
    
    # ==========================================
    # LETTERHEAD (Left aligned like CV)
    # ==========================================
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_header.paragraph_format.space_after = Pt(20)
    p_header.paragraph_format.keep_together = True
    
    run_name = p_header.add_run(clean_markdown(basics.get("name", "")) + "\n")
    run_name.bold = True
    run_name.font.size = Pt(16)
    run_name.font.color.rgb = RGBColor(0, 51, 102)
    
    location = basics.get('location', {})
    if isinstance(location, str): location = {'city': location}
    elif not isinstance(location, dict): location = {}
    city_str = f"{clean_markdown(location.get('postalCode', ''))} {clean_markdown(location.get('city', ''))}"
    address = location.get('address', '')
    if address: city_str += f" ({clean_markdown(address)})"
    
    contact_parts = []
    if basics.get('email'): contact_parts.append(clean_markdown(basics.get('email')))
    if basics.get('phone'): contact_parts.append(clean_markdown(basics.get('phone')))
    if city_str.strip(): contact_parts.append(city_str.strip())
    
    contact_text = " | ".join(contact_parts)
    if contact_text:
        p_header.add_run(contact_text).font.size = Pt(9.5)
    
    profiles = basics.get("profiles", [])
    if profiles:
        if contact_text: p_header.add_run(" | ").font.size = Pt(9.5)
        for i, prof in enumerate(profiles):
            url = prof.get("url", "") if isinstance(prof, dict) else prof
            if url:
                display = url.replace("https://", "").replace("http://", "").replace("www.", "").strip('/')
                add_hyperlink(p_header, url, display, size=9.5)
                if i < len(profiles) - 1: p_header.add_run(" | ").font.size = Pt(9.5)

    # ==========================================
    # RECIPIENT BLOCK & DATE (DIN 5008 Style)
    # ==========================================
    rec_p = doc.add_paragraph()
    rec_p.paragraph_format.space_after = Pt(15)
    
    company = clean_markdown(recipient.get('company', ''))
    person = clean_markdown(recipient.get('contact_person', ''))
    street = clean_markdown(recipient.get('address', ''))
    
    # Avoid duplicate city/postal code in rec_text
    postal = clean_markdown(recipient.get('postal_code', ''))
    city_name = clean_markdown(recipient.get('city', ''))
    city_full = f"{postal} {city_name}".strip()
    
    recipient_lines = [company, person, street, city_full]
    # Remove empty lines and duplicates (like city appearing twice)
    seen = []
    final_lines = []
    for line in recipient_lines:
        s = line.strip()
        if s and s not in seen:
            final_lines.append(s)
            seen.append(s)
            
    rec_text = "\n".join(final_lines)
    r_rec = rec_p.add_run(rec_text)
    r_rec.font.size = Pt(10)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_after = Pt(20)
    now = datetime.datetime.now()
    
    # Format: "City, DD.MM.YYYY" (No postal code in date line per request)
    city_raw = clean_markdown(letter_data.get('location', 'Essen'))
    # Clean city_raw: remove digits and leading/trailing spaces/commas
    import re
    city_only = re.sub(r'\d+', '', city_raw).strip(', ')
    
    date_str = f"{city_only}, {now.strftime('%d.%m.%Y')}"
    run_date = p_date.add_run(date_str)
    run_date.font.size = Pt(10)

    # ==========================================
    # SUBJECT LINE
    # ==========================================
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(15)
    run_sub = p_sub.add_run(clean_markdown(letter_data.get('subject', 'Bewerbung')))
    run_sub.bold = True
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(0, 51, 102)

    # ==========================================
    # SALUTATION & BODY
    # ==========================================
    salutation_text = clean_markdown(letter_data.get('salutation', 'Sehr geehrte Damen und Herren,'))
    p_sal = doc.add_paragraph(salutation_text)
    p_sal.paragraph_format.space_after = Pt(8)

    paragraphs = letter_data.get('paragraphs', [])
    for i, para in enumerate(paragraphs):
        text = clean_markdown(para)
        
        # German rule: if salutation ends with comma, first paragraph starts with lowercase (if not a noun)
        if i == 0 and salutation_text.strip().endswith(',') and text:
            # We lowercase the first letter. Note: In German, nouns are uppercase, 
            # but usually the AI starts with "ich" or "hiermit".
            text = text[0].lower() + text[1:] if len(text) > 0 else text
            
        p_body = doc.add_paragraph(text)
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_body.paragraph_format.space_after = Pt(8)
        p_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ==========================================
    # CLOSING & SIGNATURE
    # ==========================================
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p_close = doc.add_paragraph(clean_markdown(letter_data.get('closing', 'Mit freundlichen Grüßen')))
    p_close.paragraph_format.space_after = Pt(30) # Space for hand signature if printed

    p_sig = doc.add_paragraph()
    r_sig = p_sig.add_run(clean_markdown(basics.get("name", "")))
    r_sig.bold = True
    r_sig.font.size = Pt(10)

    try:
        doc.save(output_file)
        print(f"✅ Erfolg! {output_file} wurde generiert passend zum CV-Design.")
    except Exception as e:
        print(f"❌ Fehler beim Speichern: {e}")

if __name__ == "__main__":
    cv_input = "data.json"
    letter_input = "cover_letter.json"
    word_output = "Anschreiben.docx"
    if len(sys.argv) >= 2: cv_input = sys.argv[1]
    if len(sys.argv) >= 3: letter_input = sys.argv[2]
    if len(sys.argv) >= 4: word_output = sys.argv[3]
    generate_cover_letter(cv_input, letter_input, word_output)
