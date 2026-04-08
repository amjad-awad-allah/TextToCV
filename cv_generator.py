import os
import sys
import json
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn

def clean_markdown(text):
    if not text: return ""
    text = str(text)
    import re
    text = re.sub(r'\[\[([^\]]+)\]\(([^\)]+)\)\]\(([^\)]+)\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\1', text)
    text = re.sub(r'\|\s*---+\s*\|', '', text)
    text = re.sub(r'^\|\s*|\s*\|$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*|\*', '', text)
    text = re.sub(r'__|_', '', text)
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        s = line.strip()
        if not s or (s.startswith('|') and '---' in s): continue
        cleaned.append(s.lstrip('|').rstrip('|').strip())
    return '\n'.join(cleaned).strip()

def add_hyperlink(paragraph, url, text, bold=False, size=10, color='005b9f'):
    if not url:
        run = paragraph.add_run(clean_markdown(text))
        run.font.size = Pt(size)
        if bold: run.bold = True
        return
    if not url.startswith('http'):
        url = 'https://' + url
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:tooltip'), url)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Calibri')
    rFont.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFont)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.set(qn('xml:space'), 'preserve')
    text_elem.text = clean_markdown(text)
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_bullet(doc, text, font_size=10, indent=0.35, space_after=1):
    if not text: return None
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    run = p.add_run("•  ")
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run_text = p.add_run(clean_markdown(text))
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(font_size)
    return p

def format_date(date_str):
    if not date_str: return ""
    date_str = str(date_str).strip()
    if date_str.lower() in ["heute", "today", "present"]: return "heute"
    parts = date_str.split('-')
    if len(parts) == 2: return f"{parts[1]}.{parts[0]}"
    if len(parts) == 1 and len(parts[0]) == 4: return parts[0]
    return date_str

def add_section_heading(doc, title):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.keep_with_next = True
    h.paragraph_format.keep_together = True
    run = h.add_run(clean_markdown(title).upper())
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.size = Pt(11)
    run.font.bold = True
    
    separator = doc.add_paragraph()
    separator.paragraph_format.space_before = Pt(0)
    separator.paragraph_format.space_after = Pt(3)
    separator.paragraph_format.keep_with_next = True
    sep_run = separator.add_run("____________________________________________________________________________________")
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(210, 210, 210)
    return h

def add_dynamic_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    run.add_text("1")
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)


def generate_cv(data_input, output_file="Generated_CV.docx"):
    # Allow passing either a file path or a dictionary
    if isinstance(data_input, str):
        try:
            with open(data_input, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            print(f"❌ Error loading {data_input}: {e}")
            return
    else:
        data = data_input

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    
    # Page setup (DIN 5008 approx)
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(25)
        section.right_margin = Mm(20)
        
        # Footer
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run()
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(120, 120, 120)
        r.add_text("Seite ")
        add_dynamic_page_number(p)

    basics = data.get("basics", {})
    
    # HEADER - Transparent Table with Photo Support
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(5.5)
    header_table.columns[1].width = Inches(1.5)
    
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.space_after = Pt(2)
    
    # 1. Name
    run_name = p_left.add_run(clean_markdown(basics.get("name", "")) + "\n")
    run_name.bold = True
    run_name.font.size = Pt(18)
    run_name.font.color.rgb = RGBColor(0, 51, 102)
    
    # 2. Contact details + Links
    location = basics.get('location', {})
    if isinstance(location, str): location = {'city': location}
    elif not isinstance(location, dict): location = {}
    city_str = f"{clean_markdown(location.get('postalCode', ''))} {clean_markdown(location.get('city', ''))}"
    address = location.get('address', '')
    if address: city_str += f" ({clean_markdown(address)})"
    
    contact_parts = []
    if basics.get('email'): contact_parts.append(clean_markdown(basics.get('email')))
    if basics.get('phone'): contact_parts.append(clean_markdown(basics.get('phone')))
    if city_str.strip(): contact_parts.append(city_str.strip())
    
    contact_text = " | ".join(contact_parts)
    if contact_text:
        p_left.add_run(contact_text).font.size = Pt(10)
    
    profiles = basics.get("profiles", [])
    if profiles:
        if contact_text: p_left.add_run(" | ").font.size = Pt(10)
        for i, prof in enumerate(profiles):
            url = prof.get("url", "") if isinstance(prof, dict) else prof
            if url:
                display = url.replace("https://", "").replace("http://", "").replace("www.", "").strip('/')
                add_hyperlink(p_left, url, display, size=10)
                if i < len(profiles) - 1: p_left.add_run(" | ").font.size = Pt(10)
                
    # 3. Profile Photo
    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Support for photo as bytes, file path, or 'output/profile.jpg'
    photo = data.get("basics", {}).get("photo") or "output/profile.jpg"
    
    if photo:
        import io
        try:
            r_pic = p_right.add_run()
            if isinstance(photo, bytes):
                r_pic.add_picture(io.BytesIO(photo), width=Inches(1.2))
            elif os.path.exists(photo):
                r_pic.add_picture(photo, width=Inches(1.2))
        except Exception as e:
            print(f"Error loading picture: {e}")
            
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # PROFIL
    summary = basics.get("summary", "")
    if summary:
        add_section_heading(doc, 'Profil')
        ps = doc.add_paragraph(clean_markdown(summary))
        ps.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ps.paragraph_format.space_after = Pt(2)
        ps.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        ps.paragraph_format.keep_together = True
        ps.paragraph_format.widow_control = True

    # BERUFSERFAHRUNG
    work_items = data.get("work", [])
    if work_items:
        add_section_heading(doc, 'Berufserfahrung')
        for job in work_items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            
            r_left = p.add_run(f"{clean_markdown(job.get('position', ''))} | {clean_markdown(job.get('name', ''))}")
            r_left.bold = True
            r_left.font.size = Pt(10.5)
            
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Mm(165), WD_TAB_ALIGNMENT.RIGHT)
            
            date_str = f"{format_date(job.get('startDate', ''))} – {format_date(job.get('endDate', 'heute'))}"
            r_date = p.add_run(f"\t{date_str}")
            r_date.italic = True
            r_date.font.color.rgb = RGBColor(90, 90, 90)
            r_date.font.size = Pt(9.5)
            
            for bullet in job.get('highlights', []):
                bp = add_bullet(doc, bullet, font_size=9.5, indent=0.35, space_after=1)
                if bp: bp.paragraph_format.keep_with_next = False

    # AUSBILDUNG
    edu_items = data.get("education", [])
    if edu_items:
        add_section_heading(doc, 'Ausbildung')
        for edu in edu_items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            
            r = p.add_run(f"{clean_markdown(edu.get('area', ''))} | {clean_markdown(edu.get('institution', ''))}")
            r.bold = True
            r.font.size = Pt(10.5)
            
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Mm(165), WD_TAB_ALIGNMENT.RIGHT)
            date_str = f"{format_date(edu.get('startDate', ''))} – {format_date(edu.get('endDate', ''))}"
            r_date = p.add_run(f"\t{date_str}")
            r_date.italic = True
            r_date.font.color.rgb = RGBColor(90, 90, 90)
            r_date.font.size = Pt(9.5)
            
            if edu.get('notes'):
                bp = add_bullet(doc, edu.get('notes'), font_size=9.5, indent=0.35, space_after=1)
                if bp: bp.paragraph_format.keep_with_next = False

    # PROJEKTE
    proj_items = data.get("projects", [])
    if proj_items:
        add_section_heading(doc, 'Projekte')
        for proj in proj_items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            
            url = proj.get('url', '')
            name_text = f"{clean_markdown(proj.get('name', ''))}"
            
            if url:
                add_hyperlink(p, url, name_text, bold=True, size=10.5)
            else:
                r_title = p.add_run(name_text)
                r_title.bold = True
                r_title.font.size = Pt(10.5)
            
            # Description underneath
            if proj.get('description'):
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.space_after = Pt(2)
                p_desc.paragraph_format.space_before = Pt(1)
                r_desc = p_desc.add_run(clean_markdown(proj.get('description', '')))
                r_desc.italic = True
                r_desc.font.size = Pt(9.5)
            
            # Highlights
            for bullet in proj.get('highlights', []):
                bp = add_bullet(doc, bullet, font_size=9.5, indent=0.35, space_after=1)
                if bp: bp.paragraph_format.keep_with_next = False

    # TECHNISCHE KENNTNISSE
    skills_items = data.get("skills", [])
    if skills_items:
        add_section_heading(doc, 'Technische Kenntnisse')
        for skill in skills_items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.keep_together = True
            
            r = p.add_run(f"{clean_markdown(skill.get('name', ''))}: ")
            r.bold = True
            r.font.size = Pt(9.5)
            keywords = ", ".join(skill.get('keywords', []))
            p.add_run(clean_markdown(keywords)).font.size = Pt(9.5)


    # ZERTIFIKATE (1 column)
    cert_items = data.get("certificates", [])
    if cert_items:
        add_section_heading(doc, 'Zertifikate & Weiterbildung')
        for c in cert_items:
            if isinstance(c, dict):
                text = f"{clean_markdown(c.get('name', ''))} – {clean_markdown(c.get('issuer', ''))}"
                if c.get('date'):
                    formatted_date = format_date(c.get('date'))
                    text += f" ({formatted_date})"
                add_bullet(doc, text, font_size=9.5, indent=0.15, space_after=1)
            else:
                add_bullet(doc, clean_markdown(str(c)), font_size=9.5, indent=0.15, space_after=1)

    # SPRACHEN (1 column)
    lang_items = data.get("languages", [])
    if lang_items:
        add_section_heading(doc, 'Sprachen')
        for i, l in enumerate(lang_items):
            if isinstance(l, dict):
                text = f"{clean_markdown(l.get('language', ''))}: {clean_markdown(l.get('fluency', ''))}"
            else:
                text = clean_markdown(str(l))
                
            p = add_bullet(doc, text, font_size=9.5, indent=0.15, space_after=1)
            if p and i < len(lang_items) - 1:
                p.paragraph_format.keep_with_next = True

    # Save
    doc.save(output_file)
    print(f"✅ Erfolg: {output_file} wurde generiert.")
    print("📄 Bitte prüfen: Keine Markdown-Symbole, ATS freundlich, Bullet Points (•)")

if __name__ == '__main__':
    json_input = "data.json"
    word_output = "Generated_CV.docx"
    if len(sys.argv) >= 2: json_input = sys.argv[1]
    if len(sys.argv) >= 3: word_output = sys.argv[2]
    generate_cv(json_input, word_output)