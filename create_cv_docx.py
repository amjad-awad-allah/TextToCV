import sys
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT

doc = Document()

# Adjust margins to fit everything neatly on the page (A4 bounds are 8.27in width)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# Table for splitting Header and Photo Placeholder
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Inches(5.2)
table.columns[1].width = Inches(2.0)

cell_left = table.cell(0, 0)
cell_right = table.cell(0, 1)

# Contact Info (Left)
p = cell_left.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
run1 = p.add_run("Amjad Awad-Allah\n")
run1.bold = True
run1.font.size = Pt(22)
run1.font.color.rgb = RGBColor(0, 51, 102)

contact_text = (
    "amjad.awadallah93@gmail.com | +49 163 9055276\n"
    "45307 Essen (Großraum Gelsenkirchen)\n"
    "linkedin.com/in/amjad-awad-allah | github.com/amjad-awad-allah\n"
    "amjadawadallah.com"
)
run2 = p.add_run(contact_text)
run2.font.size = Pt(10)

# Photo Placeholder (Right)
p_right = cell_right.paragraphs[0]
p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_right.paragraph_format.space_after = Pt(0)
# Adding text that designates exactly where the photo should go
placeholder = p_right.add_run("\n\n[ Platzhalter für\n Bewerbungsfoto ]")
placeholder.font.color.rgb = RGBColor(120, 120, 120)
placeholder.font.italic = True

# Divider Line
p_hr = doc.add_paragraph()
p_hr.add_run("_" * 65).font.color.rgb = RGBColor(200, 200, 200)
p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_hr.paragraph_format.space_after = Pt(2)

def add_heading(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.size = Pt(12)
    run.font.bold = True
    return h

add_heading('PROFIL')
p = doc.add_paragraph("Android-Entwickler mit mehrjähriger Erfahrung in der Entwicklung von Mobile- und Backend-Lösungen für internationale Kunden. Spezialisiert auf Java, Kotlin und REST-APIs mit Fokus auf wartbare Architekturen und Clean Code. Aktuell in der Umschulung zum Fachinformatiker für Anwendungsentwicklung. Suche ein 11-monatiges Pflichtpraktikum ab Juli 2026 im Bereich Softwareentwicklung.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(2)
p.runs[0].font.size = Pt(10)

add_heading('BERUFSERFAHRUNG')

def add_job(title, company, date, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    
    # Add a right-aligned tab stop for the dates (Total width allowed ~ 7.07 inches)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    
    r = p.add_run(f"{title} | {company}")
    r.bold = True
    r.font.size = Pt(10)
    
    r2 = p.add_run(f"\t{date}")
    r2.italic = True
    r2.font.color.rgb = RGBColor(100, 100, 100)
    r2.font.size = Pt(9)
    
    for bullet in bullets:
        bp = doc.add_paragraph(bullet, style='List Bullet')
        bp.paragraph_format.space_after = Pt(1)
        bp.runs[0].font.size = Pt(10)

add_job(
    "Freelance Software-Entwickler (Teilzeit)", "Remote", "07/2021 – 05/2025",
    [
        "Entwicklung und Wartung von Android-Anwendungen für Kunden im Nahen Osten.",
        "Integration von REST-APIs und kontinuierliche Optimierung der App-Performance.",
        "Anwendung von Clean-Code-Prinzipien und strukturierten Code-Reviews zur Verbesserung der Code-Qualität.",
        "Parallele Weiterbildung: Umschulung zum Fachinformatiker + Deutschkurse (B1 → B2)."
    ]
)

add_job(
    "Android-Entwickler", "NVS-SOFT, Dubai (VAE)", "01/2018 – 09/2021",
    [
        "Entwicklung von Android-Anwendungen für Enterprise-Kunden in VAE und Kuwait.",
        "Betreuung von Softwarelösungen für staatliche Institutionen und große Unternehmen.",
        "Integration und Optimierung von REST-APIs für stabile Datenkommunikation.",
        "Zusammenarbeit mit cross-funktionalen Teams im agilen Umfeld."
    ]
)

add_job(
    "Android-Entwickler (Remote, Teilzeit)", "SmartAngle, Irak", "05/2020 – 03/2021",
    [
        "Umsetzung von Mobile-Lösungen im E-Learning- und Social-Bereich mit internationalen Teams.",
        "Anforderungsanalyse und kundenspezifische Implementierung von Features.",
        "Enge Abstimmung mit Kunden zur präzisen Umsetzung technischer Spezifikationen."
    ]
)

add_job(
    "Software-Entwickler", "PRO-Tech Group, Syrien", "02/2017 – 06/2018",
    [
        "Entwicklung von Mobile- und IPTV-Lösungen (TV-Box) für Internet- und Energieunternehmen.",
        "Integration von RESTful APIs und Streaming-Protokollen mit Fokus auf Stabilität.",
        "Erstellung technischer Dokumentation und Anwendung von Clean-Code-Standards."
    ]
)

add_heading('AUSBILDUNG')
add_job("Umschulung zum Fachinformatiker", "BFZ Essen", "2025 – heute", 
        ["Schwerpunkt: Objektorientierte Programmierung, Datenbanken, Software-Engineering, Agile Methoden, moderne Backend-Technologien (in Ausbildung)."])

add_job("B.Sc. Künstliche Intelligenz", "AIU, Damaskus", "2012–2017",
        ["Relevante Module: Algorithmen, Datenstrukturen, Grundlagen Machine Learning."])

add_heading('TECHNISCHE KENNTNISSE')
def add_skill(category, details):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(f"• {category}: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(details)
    r2.font.size = Pt(10)

add_skill("Programmiersprachen", "Sehr gut: Java, Kotlin | Gute Kenntnisse: SQL, JavaScript | Grundkenntnisse: Python, C#, Spring Boot (in Ausbildung)")
add_skill("Android-Entwicklung", "Android SDK, MVVM, MVP, Firebase, Android Studio, Jetpack Components")
add_skill("Backend & APIs", "REST API Design, JSON, HTTP, Grundlagen Backend-Entwicklung")
add_skill("Tools & Methoden", "Git, CI/CD (GitLab CI), Agile/Scrum, Clean Code, Unit Testing (Grundlagen)")
add_skill("Web & Sonstiges", "HTML5, CSS3, Responsive Design, Grundkenntnisse Unity")

add_heading('PROJEKTE (AUSWAHL)')
def add_project(name, tech, role, result):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.size = Pt(10)
    
    bullets = [f"Tech: {tech}", f"Rolle: {role}", f"Ergebnis: {result}"]
    for b in bullets:
        bp = doc.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(0)
        bp.runs[0].font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_project("Tarasol – Enterprise Document Management (2022)", "Java, Kotlin, Android SDK, REST API, SQLite, MVVM", "Android-Entwicklung mit Fokus auf Offline-Fähigkeit, Push-Benachrichtigungen und Task-Management.", "Unterstützung papierloser Workflows für aktive Nutzer im Unternehmensumfeld.")
add_project("Arcmate 9 – Enterprise Repository Client (2021)", "Java, Android SDK, Enterprise Integration, REST API", "Implementierung der Suchfunktion und Offline-Sync für Dokumenten-Repositories.", "Nahtlose Integration in bestehende Enterprise-Infrastruktur.")
add_project("Fliesen Demirel – Responsive Website (2025)", "HTML, CSS, JavaScript, Responsive Web Design", "Umsetzung eines benutzerfreundlichen Designs für Desktop, Tablet und Mobile.", "Professionelle Webpräsenz zur Präsentation von Referenzprojekten.")
add_project("Road Shield Solutions – Firmenwebsite (seit 2026)", "HTML, CSS, JavaScript, Responsive Design, Git", "Frontend-Entwicklung und Content-Management für ein Bauunternehmen in Dubai.", "In Entwicklung – Fokus auf UX und responsive Darstellung.")

add_heading('ZERTIFIKATE & WEITERBILDUNG')
def add_cert(cert):
    bp = doc.add_paragraph(cert, style='List Bullet')
    bp.paragraph_format.space_after = Pt(1)
    bp.runs[0].font.size = Pt(10)

add_cert("Cisco Networking Academy: Cybersecurity Analyst, Python, JavaScript, IoT, Grundlagen Künstliche Intelligenz")
add_cert("Deutsch: B2 – Weststadt Akademie (2025) | B1 – DÜS Eckert (2024)")
add_cert("Jobcoaching & Bewerbungstraining – Weststadt Akademie (2024–2025)")

add_heading('SPRACHEN')
add_cert("Arabisch: Muttersprache")
add_cert("Deutsch: B2 (Fortgeschritten)")
add_cert("Englisch: Fortgeschritten (Berufssprache)")
add_cert("Französisch: Grundkenntnisse")

try:
    doc.save("Amjad_Awad_CV_Professional.docx")
except Exception as e:
    print(f"Error: {e}")
