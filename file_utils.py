import docx
import PyPDF2
import io

def extract_text_from_file(uploaded_file):
    """
    Extracts text from PDF, DOCX, or TXT files.
    """
    filename = uploaded_file.name.lower()
    
    if filename.endswith(".pdf"):
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return f"Error reading PDF: {e}"
            
    elif filename.endswith(".docx"):
        try:
            doc = docx.Document(uploaded_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            return f"Error reading DOCX: {e}"
            
    elif filename.endswith(".txt"):
        try:
            return uploaded_file.read().decode("utf-8")
        except Exception as e:
            return f"Error reading TXT: {e}"
            
    return "Unsupported file format."

def extract_text_with_links_from_local(file_path):
    """
    Extracts text from a local DOCX file and preserves hyperlinks in the text.
    """
    import os
    filename = os.path.basename(file_path).lower()
    
    if filename.endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            text_lines = []
            for p in doc.paragraphs:
                text = ""
                for child in p._element:
                    if child.tag.endswith('hyperlink'):
                        rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rId and rId in p.part.rels:
                            url = p.part.rels[rId]._target
                            link_text = "".join(node.text for node in child.iter() if node.tag.endswith('t') and node.text)
                            text += f" {link_text} ({url}) "
                    elif child.tag.endswith('r'):
                        text += "".join(node.text for node in child.iter() if node.tag.endswith('t') and node.text)
                if text.strip():
                    text_lines.append(text.strip())
            return "\n".join(text_lines)
        except Exception as e:
            return f"Error reading local DOCX: {e}"
            
    return "Unsupported local file format."
